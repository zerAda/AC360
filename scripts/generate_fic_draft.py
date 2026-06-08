import os
import sys
import json
import argparse
from datetime import datetime

import csv

# Import du gestionnaire de base de données
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
try:
    from db_manager import log_fic_generation
except ImportError:
    print("Erreur : le module db_manager est introuvable.")
    exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Erreur : la librairie python-docx n'est pas installée. (pip install python-docx)")
    exit(1)

def evaluate_fic_rules(motif):
    """
    Évalue les règles métier d'Adel pour déterminer si une FIC doit être générée.
    """
    motif_lower = motif.lower()
    
    # [PATCH HATER] Critères d'exclusion de FIC (DOIVENT ÊTRE ÉVALUÉS EN PREMIER)
    if any(keyword in motif_lower for keyword in ["reprise de gestion", "changement tarif", "gestionnaire"]):
        return False, "Non requis (Reprise/Tarif/Gestionnaire)"
        
    # Critères de création de FIC
    if any(keyword in motif_lower for keyword in ["conseil", "modif", "modification", "catégorie", "categorie"]):
        return True, "Requis (Modification/Conseil)"
    
    # Par défaut, si doute, on génère
    return True, "Requis (Par défaut)"

def update_fic_tracking(client_name, motif, is_eligible, status_text):
    """
    Met à jour le tableau de bord de suivi des FIC (Base de Données SQLite).
    """
    log_fic_generation(client_name, motif, is_eligible, status_text)


def generate_fic_document(client_name, date_effet, plafonds_garanties, output_path):
    """
    Génère un document Word (.docx) contenant le brouillon de la Fiche d'Information et de Conseil (FIC).
    """
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('FICHE D\'INFORMATION ET DE CONSEIL (FIC)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # En-tête Date et Client
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_heading('1. Informations Client', level=1)
    
    p_client = doc.add_paragraph()
    p_client.add_run('Raison Sociale / Nom du Client : ').bold = True
    p_client.add_run(client_name)
    
    p_date = doc.add_paragraph()
    p_date.add_run('Date d\'effet présumée : ').bold = True
    p_date.add_run(date_effet if date_effet else "À compléter")

    # Section Besoins et Garanties
    doc.add_heading('2. Recueil des Besoins et Exigences', level=1)
    doc.add_paragraph(
        "Sur la base de l'audit de vos contrats existants et de notre échange, "
        "nous avons identifié les garanties et plafonds suivants comme base de recommandation :"
    )
    
    # Ajout d'un tableau récapitulatif des garanties trouvées
    if plafonds_garanties:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type de Garantie'
        hdr_cells[1].text = 'Valeur Actuelle (Plafond)'
        
        for garantie, valeur in plafonds_garanties.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(garantie)
            row_cells[1].text = str(valeur)
    else:
        doc.add_paragraph("Aucune donnée de garantie spécifique remontée par l'audit. [À COMPLÉTER MANUELLEMENT]")

    # Mention Devoir de Conseil
    doc.add_heading('3. Notre Conseil et Motivation', level=1)
    doc.add_paragraph(
        "Au regard de l'analyse effectuée, nous vous conseillons le maintien de vos couvertures actuelles "
        "tout en procédant à une mise en conformité de vos actes fondateurs."
    )
    
    mention = doc.add_paragraph()
    mention.add_run(
        "Conformément à la réglementation sur le devoir de conseil (Directive sur la Distribution d'Assurances - DDA), "
        "nous confirmons que cette proposition est cohérente avec vos exigences et besoins."
    )
    mention.italic = True
    
    # Signatures
    doc.add_heading('4. Signatures', level=1)
    doc.add_paragraph("Fait à ______________________, le ___/___/20__")
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.cell(0, 0).text = "Pour le Cabinet (Le Conseiller)"
    sig_table.cell(0, 1).text = "Pour le Client (Le Représentant Légal)"
    sig_table.cell(1, 0).text = "\n\n\n(Signature)"
    sig_table.cell(1, 1).text = "\n\n\n(Signature et Cachet)"

    # Enregistrement
    doc.save(output_path)
    print(f"[FIC GENERATION] Brouillon de la FIC généré avec succès : {output_path}")
    
    # Impression stricte pour récupération par PowerShell
    print(f"FIC_GENERATED_PATH={os.path.abspath(output_path)}")


def main():
    parser = argparse.ArgumentParser(description="Génère un brouillon FIC (Word) à partir des résultats d'audit.")
    parser.add_argument("audit_report", help="Fichier JSON d'audit généré à la Phase 4.")
    parser.add_argument("--output-dir", help="Dossier de sortie du brouillon.", default=".")
    args = parser.parse_args()

    if not os.path.exists(args.audit_report):
        print(f"[ERREUR] Le rapport {args.audit_report} est introuvable.")
        exit(1)

    with open(args.audit_report, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Extraction sécurisée des données
    client_name = data.get("meilleur_match_fabric") or data.get("client_document") or "CLIENT_INCONNU"
    motif_operation = data.get("motif_operation", "inconnu")
    
    print(f"Évaluation des règles FIC pour le client '{client_name}' (Motif: {motif_operation})")
    
    # Validation des règles métier d'Adel
    is_eligible, status_text = evaluate_fic_rules(motif_operation)
    
    # Mise à jour du tableau de suivi (toujours exécuté, pour la traçabilité)
    update_fic_tracking(client_name, motif_operation, is_eligible, status_text)
    
    if not is_eligible:
        print(f"[SKIP FIC] Le motif '{motif_operation}' ne requiert pas de création de FIC.")
        return

    # Construction d'un mini-dictionnaire des garanties pour le template
    garanties = {}
    for ecart in data.get("details_ecarts", []):
        champ = ecart.get("champ")
        val = ecart.get("valeur_gestion_artus") or ecart.get("valeur_document")
        if champ and val:
            garanties[champ] = val

    # Nettoyage du nom pour le fichier
    safe_client_name = "".join([c for c in client_name if c.isalpha() or c.isdigit() or c==' ']).rstrip().replace(" ", "_")
    output_filename = f"FIC_Brouillon_{safe_client_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    output_path = os.path.join(args.output_dir, output_filename)

    print(f"Génération du document pour {client_name}...")
    generate_fic_document(client_name, "", garanties, output_path)

if __name__ == "__main__":
    main()
