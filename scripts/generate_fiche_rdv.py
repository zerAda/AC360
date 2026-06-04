import os
import re
import uuid
import unicodedata
import datetime
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from config import JOBS_BASE_DIR
except ImportError:
    JOBS_BASE_DIR = os.path.abspath("jobs")


def safe_filename(name: str, max_length: int = 64) -> str:
    """
    Sanitise un nom de fichier pour prévenir le Path Traversal (P0-05).
    - Normalise les accents (NFD -> ASCII)
    - Supprime tous les caractères hors [a-zA-Z0-9_-]
    - Interdit les séquences '..'
    - Limite la longueur
    """
    # Normaliser les accents
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
    # Allowlist stricte
    name = re.sub(r"[^a-zA-Z0-9 _\-]", "", name)
    # Remplacer espaces par underscore
    name = name.strip().replace(" ", "_")
    # Interdire '..' explicitement
    name = name.replace("..", "")
    # Limiter la longueur
    name = name[:max_length] if len(name) > max_length else name
    # Fallback si nom vide après sanitization
    if not name:
        name = "client_inconnu"
    return name


def generate_fiche_rdv(client_name: str, summary: str, alert_points: str, job_id: str = None) -> str:
    """
    Génère un document Word synthétique pour préparer un rendez-vous.
    Protégé contre le Path Traversal via safe_filename().
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "Le module 'python-docx' n'est pas installé. "
            "Exécutez: pip install python-docx"
        )

    if not job_id:
        job_id = str(uuid.uuid4())

    # Sanitisation du nom client avant usage dans le chemin
    safe_name = safe_filename(client_name)

    job_dir = Path(JOBS_BASE_DIR) / job_id
    job_dir.mkdir(parents=True, exist_ok=True)

    file_path = job_dir / f"Fiche_RDV_{safe_name}.docx"

    # Vérification Path Traversal via commonpath
    resolved = file_path.resolve()
    base_resolved = Path(JOBS_BASE_DIR).resolve()
    if not str(resolved).startswith(str(base_resolved)):
        raise PermissionError(f"Path traversal détecté : {resolved}")

    doc = Document()
    doc.add_heading(f"Fiche de Rendez-vous : {client_name}", 0)

    p_date = doc.add_paragraph()
    p_date.add_run(f"Généré le : {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}").italic = True

    doc.add_heading("1. Synthèse du Dossier", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("2. Points d'Attention & Alertes", level=1)
    doc.add_paragraph(alert_points)

    doc.add_heading("3. Notes du Commercial", level=1)
    doc.add_paragraph("\n\n\n[Espace pour prendre des notes pendant le RDV...]")

    doc.save(str(file_path))

    return str(file_path)


if __name__ == "__main__":
    path = generate_fiche_rdv(
        "Client Alpha",
        "Contrat mutuelle standard. 50 employés.",
        "Kbis manquant. Renouvellement dans 2 mois."
    )
    print(f"Fiche générée : {path}")
