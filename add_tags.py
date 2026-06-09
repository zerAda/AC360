import docx

def add_tags_to_docx(input_path, output_path):
    doc = docx.Document(input_path)
    
    replacements = {
        "Raison Sociale :": "Raison Sociale : {{ raison_sociale }}",
        "Forme Juridique\xa0:": "Forme Juridique : {{ forme_juridique }}",
        "Numéro SIREN/SIRET :": "Numéro SIREN/SIRET : {{ siren }}",
        "Secteur d’activité principale :": "Secteur d’activité principale : {{ secteur_activite }}",
        "Adresse\xa0:": "Adresse : {{ adresse }}",
        "Convention Collective de rattachement :": "Convention Collective de rattachement : {{ ccn }}",
        "lors de notre rendez-vous du XX XX XX": "lors de notre rendez-vous du {{ date_rdv }}",
        "Fait à Paris, le XX XX 2025": "Fait à Paris, le {{ date_jour }}"
    }
    
    # Remplacements alternatifs (problèmes d'espaces insécables)
    alt_replacements = {
        "Forme Juridique :": "Forme Juridique : {{ forme_juridique }}",
        "Adresse :": "Adresse : {{ adresse }}",
        "Secteur d'activité principale :": "Secteur d’activité principale : {{ secteur_activite }}"
    }

    # Pour éviter les problèmes de text runs, on efface tous les runs d'un paragraphe et on recrée le texte
    # si le paragraphe contient la clé (car le texte est simple).
    for p in doc.paragraphs:
        original_text = p.text
        if not original_text.strip():
            continue
            
        modified = False
        new_text = original_text
        
        for k, v in replacements.items():
            if k in new_text:
                new_text = new_text.replace(k, v)
                modified = True

        if modified:
            # On conserve le style du premier run si possible
            style = p.style
            for r in p.runs:
                r.text = ""
            p.add_run(new_text)

    doc.save(output_path)
    print(f"Modèle sauvegardé avec succès vers {output_path}")

if __name__ == '__main__':
    add_tags_to_docx('FIC_Modèle.docx', 'FIC_Modèle_Tagge.docx')
